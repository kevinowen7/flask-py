from flask import Flask, request, jsonify

import docx
import os
from io import BytesIO
from datetime import datetime
from PIL import Image
import pytesseract
import aspose.words as aw

import firebase_admin
from firebase_admin import credentials
from firebase_admin import storage

import io

app = Flask(__name__)

#pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe'

cred = credentials.Certificate('serviceAccountKey.json')
firebase_admin.initialize_app(cred, {
    'storageBucket': 'project-ta-4e64c.appspot.com'
})
bucket = storage.bucket()

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

@app.route('/')
def hello_world():
    return 'Hello, World!'

@app.route('/convert-image', methods=['POST'])
def convert_image():
    if 'file' not in request.files:
        return jsonify(
            success=0,
            message="File not found",
            data=None
        )
    
    image = request.files['file']

    # upload origin file to firebase
    filename_original = datetime.now().strftime('%Y%m%d%H%M%S')+".jpg"

    f = io.BytesIO()
    pil_img = Image.open(image)
    pil_img.save(f, format='JPEG')
    pil_img.close()
    
    blob_original = bucket.blob("origin/"+filename_original)
    blob_original.upload_from_string(f.getvalue())
    blob_original.make_public()
    
    # OCR
    text_parsed = pytesseract.image_to_string(Image.open(image))
    
    # Create a document
    doc = docx.Document()

    # Add a paragraph to the document
    p = doc.add_paragraph()
    p.add_run(''.join(c for c in text_parsed if valid_xml_char_ordinal(c)))

    # save document
    filename = datetime.now().strftime('%Y%m%d%H%M%S')+".docx"
    doc.save(filename)
    
    streamFile = io.FileIO(filename)
    awDocFile = aw.Document(streamFile)
    streamFile.close()

    extractedPageName = datetime.now().strftime('%Y%m%d%H%M%S')+".jpg"
    for page in range(0, awDocFile.page_count):
        extractedPage = awDocFile.extract_pages(page, 1)
        extractedPage.watermark.remove()
        extractedPage.save(extractedPageName)

    # upload to firebase
    blob = bucket.blob("result/"+filename)
    blob.upload_from_filename(filename)
    blob.make_public()
    os.unlink(filename)

    # upload to firebase
    blob_preview = bucket.blob("preview/"+extractedPageName)
    blob_preview.upload_from_filename(extractedPageName)
    blob_preview.make_public()
    os.unlink(extractedPageName)

    resultData = {}
    resultData["result_file"] = blob.public_url
    resultData["original_file"] = blob_original.public_url
    resultData["preview_file"] = blob_preview.public_url
    
    # END
    return jsonify(
        success=1,
        message="Success",
        data=resultData
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)
